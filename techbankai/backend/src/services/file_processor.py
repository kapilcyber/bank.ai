"""File processing utilities for extracting text from various file formats."""
import PyPDF2
import pdfplumber
from docx import Document
from typing import Optional
from src.utils.logger import get_logger
import shutil
import subprocess
import tempfile
import os

logger = get_logger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file using pdfplumber (more accurate).
    Falls back to PyPDF2 if pdfplumber fails.
    """
    try:
        # Try pdfplumber first (better for complex PDFs)
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            if text.strip():
                logger.info(f"Extracted {len(text)} characters from PDF using pdfplumber")
                return text.strip()
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")
    
    # Fallback to PyPDF2
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            logger.info(f"Extracted {len(text)} characters from PDF using PyPDF2")
            return text.strip()
    except Exception as e:
        logger.error(f"Failed to extract PDF text: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file (paragraphs + tables + headers/footers)."""
    try:
        doc = Document(file_path)

        chunks = []

        # Paragraphs
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                chunks.append(p.text.strip())

        # Tables (many resumes are table-based)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text and p.text.strip():
                            chunks.append(p.text.strip())

        # Headers/footers (contact info sometimes lives here)
        try:
            for section in doc.sections:
                for p in section.header.paragraphs:
                    if p.text and p.text.strip():
                        chunks.append(p.text.strip())
                for p in section.footer.paragraphs:
                    if p.text and p.text.strip():
                        chunks.append(p.text.strip())
        except Exception:
            pass

        text = "\n".join(chunks).strip()
        logger.info(f"Extracted {len(text)} characters from DOCX")
        return text
    except Exception as e:
        logger.error(f"Failed to extract DOCX text: {e}")
        return ""


def extract_text_from_doc(file_path: str) -> str:
    """
    Extract text from DOC file (old format).
    Strategy (best-effort):
    - Windows + MS Word installed: convert to DOCX via COM automation (pywin32), then parse DOCX
    - LibreOffice (soffice) available: convert to DOCX, then parse DOCX
    - antiword available: extract plain text directly
    """
    # 1) Try MS Word COM automation on Windows (requires pywin32 + Word installed)
    try:
        import win32com.client  # type: ignore

        tmp_dir = tempfile.mkdtemp(prefix="doc_convert_")
        out_docx = os.path.join(tmp_dir, f"{os.path.splitext(os.path.basename(file_path))[0]}.docx")

        word = None
        doc = None
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(file_path, ReadOnly=True)
            # 16 = wdFormatDocumentDefault (DOCX)
            doc.SaveAs(out_docx, FileFormat=16)
        finally:
            try:
                if doc is not None:
                    doc.Close(False)
            except Exception:
                pass
            try:
                if word is not None:
                    word.Quit()
            except Exception:
                pass

        if os.path.exists(out_docx):
            return extract_text_from_docx(out_docx)
    except Exception as e:
        logger.info(f"DOC->DOCX via Word COM not available/failed: {e}")

    # 2) Try LibreOffice conversion (soffice) if installed
    soffice = shutil.which("soffice") or shutil.which("soffice.exe")
    if soffice:
        try:
            tmp_dir = tempfile.mkdtemp(prefix="doc_convert_")
            # --headless conversion writes output to --outdir with same basename
            subprocess.run(
                [soffice, "--headless", "--convert-to", "docx", "--outdir", tmp_dir, file_path],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=120,
            )
            base = os.path.splitext(os.path.basename(file_path))[0]
            out_docx = os.path.join(tmp_dir, f"{base}.docx")
            if os.path.exists(out_docx):
                return extract_text_from_docx(out_docx)
        except Exception as e:
            logger.info(f"DOC->DOCX via LibreOffice not available/failed: {e}")

    # 3) Try antiword (plain text extractor for .doc)
    antiword = shutil.which("antiword") or shutil.which("antiword.exe")
    if antiword:
        try:
            proc = subprocess.run(
                [antiword, file_path],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=60,
            )
            text = proc.stdout.decode(errors="ignore").strip()
            if text:
                logger.info(f"Extracted {len(text)} characters from DOC using antiword")
                return text
        except Exception as e:
            logger.info(f"DOC text extraction via antiword failed: {e}")

    logger.error(
        "DOC resume received but no DOC extractor is available. "
        "Install Microsoft Word (Windows) + pywin32, or LibreOffice (soffice), or antiword; "
        "or ask candidates to send PDF/DOCX."
    )
    return ""


def extract_text_from_file(file_path: str, file_extension: str) -> str:
    """Extract text from file based on extension."""
    ext = file_extension.lower().replace('.', '')
    
    if ext == 'pdf':
        return extract_text_from_pdf(file_path)
    elif ext == 'docx':
        return extract_text_from_docx(file_path)
    elif ext == 'doc':
        return extract_text_from_doc(file_path)
    else:
        logger.error(f"Unsupported file extension: {ext}")
        return ""

